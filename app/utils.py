# Helpers: text extraction, chunking

import io
import os
from typing import List, Union
from pypdf import PdfReader
import docx

def extract_text(file_input: Union[str, bytes], file_name: str = None) -> str:
    if isinstance(file_input, bytes):
        if not file_name:
            raise ValueError("File name is required for byte-based extraction")
        file_name = file_name.lower()
        if file_name.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(file_input))
            return " ".join([page.extract_text() or "" for page in reader.pages])
        elif file_name.endswith(".docx"):
            doc = docx.Document(io.BytesIO(file_input))
            return " ".join([para.text for para in doc.paragraphs])
        elif file_name.endswith(".txt"):
            return file_input.decode("utf-8", errors="ignore")
        else:
            raise ValueError("Unsupported file format")

    if isinstance(file_input, str):
        if file_input.endswith(".pdf"):
            reader = PdfReader(file_input)
            return " ".join([page.extract_text() or "" for page in reader.pages])
        elif file_input.endswith(".docx"):
            doc = docx.Document(file_input)
            return " ".join([para.text for para in doc.paragraphs])
        elif file_input.endswith(".txt"):
            with open(file_input, "r", encoding="utf-8", errors="ignore") as fh:
                return fh.read()
        else:
            raise ValueError("Unsupported file format")

    raise ValueError("Unsupported input type for text extraction")